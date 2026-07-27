#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
render.py - coule un courrier dans le modele officiel BENSAID AVOCATS.

Usage :
    python3 .claude/skills/courrier/render.py /chemin/spec.json /chemin/sortie.docx

Le script part de assets/modele_courrier.docx (en-tete a logo + pied de page
multi-sites deja en place) et ne reconstruit QUE le corps du courrier, en
respectant la charte du modele (Arial, couleur 1E1E22, N/Ref. et P.J. en 9pt,
signature alignee a droite). Il ne touche ni a l'en-tete ni au pied de page.

------------------------------------------------------------------------------
MODELE DE SPEC JSON
------------------------------------------------------------------------------
{
  "ville": "Paris",                       # defaut : Paris
  "date": "8 juillet 2026",               # obligatoire (texte libre, deja en francais)
  "destinataire": [                        # liste de lignes ; la 1re ligne est en gras
    "Monsieur Jean DUPONT",
    "Gerant",
    "12 rue de la Paix",
    "75002 Paris"
  ],
  "ref": "2026-0142 / JD",                 # optionnel : si absent, la ligne N/Ref. disparait
  "objet": "Cession de parts sociales",    # obligatoire (affiche en gras)
  "salutation": "Monsieur,",               # obligatoire
  "corps": [                                # obligatoire : liste de blocs
    "Premier paragraphe en clair.",         #   string  -> paragraphe justifie
    {"p": "Autre paragraphe."},             #   {"p":...} -> paragraphe justifie
    {"sub": "Un sous-titre en gras"},       #   {"sub":...} -> sous-titre gras
    {"b": "Une puce"},                       #   {"b":...} -> puce
    {"em": "Note en italique gris"}          #   {"em":...} -> italique gris
  ],
  "politesse": "Je vous prie d'agreer, Monsieur, l'expression de mes salutations distinguees.",
  "signataire": "Francois OUAIRY",         # obligatoire
  "qualite": "Avocat associe",             # optionnel
  "signature": "assets/signature_ouairy.png", # optionnel : image manuscrite au-dessus du nom ;
                                           #   defaut = signature Francois ; false = aucune
  "signature_largeur_cm": 4.5,             # optionnel : largeur de l'image de signature
  "pj": ["Projet d'acte de cession", "RIB"] # optionnel : liste ; si absent, ligne P.J. supprimee
}
------------------------------------------------------------------------------
"""

import sys
import json
import copy
import os

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm

HERE = os.path.dirname(os.path.abspath(__file__))
TEMPLATE = os.path.join(HERE, "assets", "modele_courrier.docx")

GREY = "808088"   # gris pour les notes italiques (charte cabinet)


# ---------------------------------------------------------------------------
# helpers bas niveau sur les paragraphes / runs
# ---------------------------------------------------------------------------

def para_text(p):
    return "".join(node.text or "" for node in p._p.iter(qn("w:t")))


def find_para(doc, marker):
    """Retourne le premier paragraphe du corps contenant `marker`."""
    for p in doc.paragraphs:
        if marker in para_text(p):
            return p
    return None


def first_run_rpr(p):
    """Renvoie une copie du <w:rPr> du premier run (pour conserver la mise en forme)."""
    r = p._p.find(qn("w:r"))
    if r is None:
        return None
    rpr = r.find(qn("w:rPr"))
    return copy.deepcopy(rpr) if rpr is not None else None


def clear_runs(p):
    """Supprime tous les runs du paragraphe, garde le pPr."""
    for r in p._p.findall(qn("w:r")):
        p._p.remove(r)


def set_text_keep_format(p, text, rpr=None):
    """Vide le paragraphe et y ecrit `text`, en conservant le rPr fourni (ou celui du 1er run)."""
    if rpr is None:
        rpr = first_run_rpr(p)
    clear_runs(p)
    run = p._p.makeelement(qn("w:r"), {})
    if rpr is not None:
        run.append(copy.deepcopy(rpr))
    t = p._p.makeelement(qn("w:t"), {})
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    p._p.append(run)


def delete_para(p):
    p._p.getparent().remove(p._p)


def clone_after(p):
    """Clone le paragraphe `p` (pPr + rPr conserves) et l'insere juste apres. Renvoie le clone."""
    new_p = copy.deepcopy(p._p)
    p._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, p._parent)


def insert_image_before(p, img_path, width_cm=4.5, align=None):
    """Insere un paragraphe contenant l'image `img_path` juste AVANT le paragraphe `p`."""
    from docx.text.paragraph import Paragraph
    new_el = p._p.makeelement(qn("w:p"), {})
    p._p.addprevious(new_el)
    new_p = Paragraph(new_el, p._parent)
    new_p.add_run().add_picture(img_path, width=Cm(width_cm))
    if align is not None:
        new_p.alignment = align
    return new_p


def set_bold(p, value=True):
    for r in p._p.findall(qn("w:r")):
        rpr = r.find(qn("w:rPr"))
        if rpr is None:
            rpr = r.makeelement(qn("w:rPr"), {})
            r.insert(0, rpr)
        b = rpr.find(qn("w:b"))
        if value and b is None:
            rpr.append(rpr.makeelement(qn("w:b"), {}))
        elif not value and b is not None:
            rpr.remove(b)


def set_italic_grey(p):
    for r in p._p.findall(qn("w:r")):
        rpr = r.find(qn("w:rPr"))
        if rpr is None:
            rpr = r.makeelement(qn("w:rPr"), {})
            r.insert(0, rpr)
        if rpr.find(qn("w:i")) is None:
            rpr.append(rpr.makeelement(qn("w:i"), {}))
        col = rpr.find(qn("w:color"))
        if col is None:
            col = rpr.makeelement(qn("w:color"), {})
            rpr.append(col)
        col.set(qn("w:val"), GREY)


# ---------------------------------------------------------------------------
# construction du courrier
# ---------------------------------------------------------------------------

def build(spec, out_path):
    doc = Document(TEMPLATE)

    # --- Date (aligne a droite dans le modele) ---
    p_date = find_para(doc, "[Date]")
    ville = spec.get("ville", "Paris")
    set_text_keep_format(p_date, f"{ville}, le {spec['date']}")

    # --- Destinataire : 4 paragraphes placeholder (Nom gras, Qualite, Adresse, CP Ville) ---
    p_nom = find_para(doc, "[NOM DU DESTINATAIRE]")
    p_qual = find_para(doc, "[QUALITE / FONCTION]")
    p_adr = find_para(doc, "[ADRESSE LIGNE 1]")
    p_cp = find_para(doc, "[CODE POSTAL, VILLE]")
    dest = spec.get("destinataire", [])
    # rPr de reference : gras pour la 1re ligne, normal pour les suivantes
    rpr_bold = first_run_rpr(p_nom)
    rpr_norm = first_run_rpr(p_adr)
    dest_slots = [p_nom, p_qual, p_adr, p_cp]

    if dest:
        # 1re ligne dans le slot "nom" (garde le gras)
        set_text_keep_format(p_nom, dest[0], rpr_bold)
        # lignes suivantes : on reutilise les slots restants puis on clone si besoin
        anchor = p_nom
        rest = dest[1:]
        reuse = [p_qual, p_adr, p_cp]
        for i, line in enumerate(rest):
            if i < len(reuse):
                target = reuse[i]
                set_text_keep_format(target, line, rpr_norm)
                anchor = target
            else:
                anchor = clone_after(anchor)
                set_text_keep_format(anchor, line, rpr_norm)
        # supprimer les slots inutilises
        for extra in reuse[len(rest):]:
            delete_para(extra)
    else:
        for p in dest_slots:
            delete_para(p)

    # --- N/Ref. ---
    p_ref = find_para(doc, "[Référence dossier]")
    if spec.get("ref"):
        set_text_keep_format(p_ref, f"N/Réf. : {spec['ref']}")
    else:
        delete_para(p_ref)

    # --- Objet (gras) ---
    p_objet = find_para(doc, "[Objet du courrier]")
    set_text_keep_format(p_objet, f"Objet : {spec['objet']}", first_run_rpr(p_objet))

    # --- Salutation ---
    p_salut = find_para(doc, "[Madame / Monsieur / Cher Maître]")
    set_text_keep_format(p_salut, spec["salutation"])

    # --- Corps ---
    p_body = find_para(doc, "[Rédigez votre courrier ici]")
    rpr_body = first_run_rpr(p_body)
    corps = spec.get("corps", [])
    if not corps:
        corps = [""]
    anchor = p_body
    for i, block in enumerate(corps):
        target = p_body if i == 0 else clone_after(anchor)
        anchor = target
        if isinstance(block, str):
            set_text_keep_format(target, block, rpr_body)
        elif "p" in block:
            set_text_keep_format(target, block["p"], rpr_body)
        elif "sub" in block:
            set_text_keep_format(target, block["sub"], rpr_body)
            set_bold(target, True)
        elif "b" in block:
            set_text_keep_format(target, "• " + block["b"], rpr_body)
        elif "em" in block:
            set_text_keep_format(target, block["em"], rpr_body)
            set_italic_grey(target)
        else:
            set_text_keep_format(target, str(block), rpr_body)

    # --- Formule de politesse ---
    p_pol = find_para(doc, "[Formule de politesse]")
    set_text_keep_format(p_pol, spec.get("politesse",
        "Je vous prie d'agréer, " + spec["salutation"].rstrip(",") +
        ", l'expression de mes salutations distinguées."))

    # --- Signature : image manuscrite au-dessus du nom, puis nom et qualite ---
    p_sig = find_para(doc, "[Prénom NOM]") or find_para(doc, "François OUAIRY")
    if p_sig is not None:
        set_text_keep_format(p_sig, spec["signataire"], first_run_rpr(p_sig))
        # image de signature au-dessus du nom : chemin fourni, sinon signature par defaut du
        # cabinet (Francois OUAIRY). Mettre "signature": false pour n'en poser aucune.
        sig = spec.get("signature", os.path.join(HERE, "assets", "signature_ouairy.png"))
        if sig and os.path.exists(sig):
            insert_image_before(p_sig, sig, spec.get("signature_largeur_cm", 4.5), p_sig.alignment)
    p_qualite = find_para(doc, "[Qualité]") or find_para(doc, "Avocat associé")
    if p_qualite is not None:
        if spec.get("qualite"):
            set_text_keep_format(p_qualite, spec["qualite"], first_run_rpr(p_qualite))
        else:
            delete_para(p_qualite)

    # --- P.J. ---
    p_pj = find_para(doc, "[Liste des pièces jointes]")
    pj = spec.get("pj")
    if pj:
        if isinstance(pj, list):
            pj = ", ".join(pj)
        set_text_keep_format(p_pj, f"P.J. : {pj}")
    else:
        delete_para(p_pj)

    doc.save(out_path)
    print("Courrier genere :", out_path)


def main():
    if len(sys.argv) != 3:
        print(__doc__)
        sys.exit(1)
    spec = json.load(open(sys.argv[1], encoding="utf-8"))
    build(spec, sys.argv[2])


if __name__ == "__main__":
    main()
