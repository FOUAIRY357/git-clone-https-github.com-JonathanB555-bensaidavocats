#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
render.py - coule un ecrit du cabinet dans un gabarit BENSAID AVOCATS.

Usage :
    python3 render.py <spec.json> [sortie.docx]

Le gabarit (en-tete a logo + pied de page coordonnees) est conserve tel quel ;
seul le corps est reconstruit selon la charte (Helvetica Neue, titres 111111,
accents 808088). Aucune couleur hors charte.

Le JSON de specification :
{
  "type": "note" | "courrier",
  "sortie": "chemin/vers/sortie.docx",          # optionnel (sinon 2e arg CLI)
  "meta": {
     # commun
     "objet": "...", "ref": "...", "signataire": "Francois OUAIRY", "qualite": "Avocat",
     # note :
     "titre": "NOTE JURIDIQUE",   # ou "COMPTE RENDU", etc.
     "confidentiel": true,
     "date": "Paris, le 24 juin 2026", "pour": "Madame X", "de": "Francois Ouairy",
     "sommaire": true,            # genere le sommaire a partir des titres de sections
     "intro": "texte d'introduction",
     "disclaimer": "texte de la mention finale (optionnel, defaut fourni)",
     # courrier :
     "ville_date": "Paris, le 24 juin 2026",
     "destinataire": ["Madame Dorothee Lalanne", "Adresse 1", "1206 Geneve"],
     "salutation": "Madame,",
     "politesse": "Je vous prie d'agreer, Madame, l'expression de mes salutations distinguees.",
     "pj": "..."                  # optionnel
  },
  "sections": [                   # note : parties numerotees ; courrier : corps
     {"titre": "Oiseau de Pauline - avis defavorable", "blocs": [ ...blocs... ]},
     ...
  ],
  "corps": [ ...blocs... ]        # courrier : alternative a sections (corps simple)
}

Types de blocs :
  {"p": "..."}            paragraphe justifie (corps)
  {"sub": "..."}          sous-titre gras
  {"b": "..."}            puce
  {"em": "..."}           note en italique gris
  {"img": ["a.jpg","b.jpg"], "w": 7.3}   grille d'images (1 = centree ; >1 = 2 colonnes)
  {"tbl": [["c1","c2"],["..."]], "entete": true}   tableau (1re ligne = entete si entete)
"""
import sys, json, os, math
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

FONT = "Helvetica Neue"
NOIR = RGBColor(0x11, 0x11, 0x11)
GRIS = RGBColor(0x80, 0x80, 0x88)
HERE = os.path.dirname(os.path.abspath(__file__))
DISCLAIMER_DEF = ("La presente note est etablie sur la base des faits et documents qui nous ont ete communiques. "
                  "Elle est confidentielle et reservee a son destinataire ; elle ne peut etre produite ou "
                  "diffusee sans notre accord.")

def _sectpr(doc):
    return doc.element.body.find(qn('w:sectPr'))

def _clear_body(doc):
    body = doc.element.body
    for child in list(body):
        if child.tag == qn('w:sectPr'):
            continue
        body.remove(child)

def _p(doc, style=None):
    """Nouveau paragraphe insere avant le sectPr (preserve en-tete/pied)."""
    body = doc.element.body
    sect = _sectpr(doc)
    el = OxmlElement('w:p')
    if sect is not None:
        sect.addprevious(el)
    else:
        body.append(el)
    par = Paragraph(el, doc._body)
    if style:
        try:
            par.style = doc.styles[style]
        except KeyError:
            pass
    return par

def _has_style(doc, name):
    try:
        doc.styles[name]; return True
    except KeyError:
        return False

def _run(par, text, size=10, bold=False, italic=False, color=None):
    r = par.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return r

def _line(doc, text, size=10, bold=False, italic=False, color=None, align=None,
          after=4, before=0, style=None):
    par = _p(doc, style=style)
    if text != "":
        _run(par, text, size, bold, italic, color)
    par.paragraph_format.space_after = Pt(after)
    par.paragraph_format.space_before = Pt(before)
    if align is not None:
        par.alignment = align
    return par

def _img_grid(doc, paths, w=None):
    paths = [p for p in paths if p and os.path.exists(p)]
    if not paths:
        return
    if len(paths) == 1:
        par = _p(doc)
        par.alignment = AL.CENTER
        par.add_run().add_picture(paths[0], width=Cm(w if w else 9))
        par.paragraph_format.space_after = Pt(6)
        return
    w = w or 7.3
    rows = math.ceil(len(paths) / 2)
    tbl = doc.add_table(rows=rows, cols=2)
    # deplacer la table avant le sectPr
    sect = _sectpr(doc)
    if sect is not None:
        sect.addprevious(tbl._tbl)
    i = 0
    for r in range(rows):
        for c in range(2):
            cell = tbl.cell(r, c)
            cp = cell.paragraphs[0]
            cp.alignment = AL.CENTER
            if i < len(paths):
                cp.add_run().add_picture(paths[i], width=Cm(w))
                i += 1
    _line(doc, "", after=6)

def _set_borders(tbl):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), '808088')
        borders.append(e)
    tblPr.append(borders)

def _table(doc, rows, entete=True):
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    _set_borders(tbl)
    sect = _sectpr(doc)
    if sect is not None:
        sect.addprevious(tbl._tbl)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cp = tbl.cell(ri, ci).paragraphs[0]
            r = cp.add_run(str(val)); r.font.name = FONT; r.font.size = Pt(9.5)
            if entete and ri == 0:
                r.bold = True
                r.font.color.rgb = NOIR
    _line(doc, "", after=6)

def _blocs(doc, blocs):
    for blk in blocs:
        if "p" in blk:
            par = _line(doc, blk["p"], size=10.5, after=12)
            par.alignment = AL.JUSTIFY
            par.paragraph_format.line_spacing = 1.4
            par.paragraph_format.first_line_indent = Cm(0.6)
        elif "sub" in blk:
            _line(doc, blk["sub"], size=10.5, bold=True, color=NOIR, after=4, before=8)
        elif "b" in blk:
            if _has_style(doc, 'List Bullet'):
                par = _line(doc, blk["b"], size=10.5, after=6, style='List Bullet')
            else:
                par = _line(doc, "•  " + blk["b"], size=10.5, after=6)
                par.paragraph_format.left_indent = Cm(0.6)
            par.paragraph_format.line_spacing = 1.3
        elif "em" in blk:
            _line(doc, blk["em"], size=9, italic=True, color=GRIS, after=8, before=4)
        elif "img" in blk:
            _img_grid(doc, blk["img"], blk.get("w"))
        elif "tbl" in blk:
            _table(doc, blk["tbl"], blk.get("entete", True))

def render_note(doc, spec):
    m = spec.get("meta", {})
    secs = spec.get("sections", [])
    if m.get("confidentiel", True):
        _line(doc, "CONFIDENTIEL", size=8, bold=True, color=GRIS, after=2)
    _line(doc, m.get("titre", "NOTE JURIDIQUE"), size=12, bold=True, color=NOIR, after=8)
    for label, key in [("Date :  ", "date"), ("N/Ref. :  ", "ref"),
                       ("Pour :  ", "pour"), ("De :  ", "de")]:
        if m.get(key):
            par = _p(doc); _run(par, label, 9, color=GRIS); _run(par, m[key], 9, color=GRIS)
            par.paragraph_format.space_after = Pt(1)
    if m.get("objet"):
        par = _p(doc); _run(par, "Objet :  ", 9, bold=True, color=GRIS); _run(par, m["objet"], 9, bold=True, color=GRIS)
        par.paragraph_format.space_after = Pt(8)
    # sommaire
    romains = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X", "XI", "XII"]
    if m.get("sommaire") and secs:
        _line(doc, "SOMMAIRE", size=9, bold=True, color=NOIR, after=6, before=4)
        for i, s in enumerate(secs):
            _line(doc, "%s.   %s" % (romains[i], s["titre"]), size=9.5, bold=True, after=3)
        _line(doc, "", after=12)
    # intro
    if m.get("intro"):
        _line(doc, "INTRODUCTION", size=10.5, bold=True, color=NOIR, after=6, before=12)
        par = _line(doc, m["intro"], size=10.5, after=14)
        par.alignment = AL.JUSTIFY
        par.paragraph_format.line_spacing = 1.4
        par.paragraph_format.first_line_indent = Cm(0.6)
    # sections numerotees
    for i, s in enumerate(secs):
        _line(doc, "%s.   %s" % (romains[i], s["titre"]), size=11, bold=True, color=NOIR, after=8, before=18)
        _blocs(doc, s.get("blocs", []))
    # disclaimer + signature
    _line(doc, m.get("disclaimer", DISCLAIMER_DEF), size=8.5, italic=True, color=GRIS, after=4, before=10)
    _line(doc, "", after=0)
    if m.get("signataire"):
        # image de signature manuscrite au-dessus du nom (comme le skill courrier) :
        # defaut = signature Francois OUAIRY ; "signature": false pour n'en poser aucune ;
        # "signature": "<chemin>" pour un autre signataire.
        sig = m.get("signature", os.path.join(HERE, "assets", "signature_ouairy.png"))
        if sig and os.path.exists(sig):
            ps = _p(doc)
            ps.add_run().add_picture(sig, width=Cm(m.get("signature_largeur_cm", 4.0)))
            ps.paragraph_format.space_after = Pt(0)
            ps.paragraph_format.space_before = Pt(6)
        _line(doc, m["signataire"], size=9.5, after=0)
    if m.get("qualite"):
        _line(doc, m["qualite"], size=9, color=GRIS, after=0)

def render_courrier(doc, spec):
    m = spec.get("meta", {})
    _line(doc, m.get("ville_date", ""), size=10, after=12)
    for i, ligne in enumerate(m.get("destinataire", [])):
        _line(doc, ligne, size=10, bold=(i == 0), after=1)
    if m.get("ref"):
        _line(doc, "N/Ref. : " + m["ref"], size=9, color=GRIS, after=8, before=8)
    if m.get("objet"):
        _line(doc, "Objet : " + m["objet"], size=10, bold=True, after=10)
    if m.get("salutation"):
        _line(doc, m["salutation"], size=10, after=8)
    blocs = spec.get("corps")
    if blocs is None:
        blocs = []
        for s in spec.get("sections", []):
            if s.get("titre"):
                blocs.append({"sub": s["titre"]})
            blocs.extend(s.get("blocs", []))
    _blocs(doc, blocs)
    if m.get("politesse"):
        par = _line(doc, m["politesse"], size=10, after=12, before=6); par.alignment = AL.JUSTIFY
    if m.get("signataire"):
        _line(doc, m["signataire"], size=9.5, after=0, before=6)
    if m.get("qualite"):
        _line(doc, m["qualite"], size=9, color=GRIS, after=0)
    if m.get("pj"):
        _line(doc, "P.J. : " + m["pj"], size=9, color=GRIS, after=0, before=10)

def render(spec, out=None):
    typ = spec.get("type", "note")
    tpl = os.path.join(HERE, "templates", "note_juridique.docx" if typ == "note" else "courrier.docx")
    doc = Document(tpl)
    _clear_body(doc)
    if typ == "note":
        render_note(doc, spec)
    else:
        render_courrier(doc, spec)
    out = out or spec.get("sortie")
    if not out:
        raise SystemExit("Chemin de sortie manquant (cle 'sortie' ou 2e argument).")
    doc.save(out)
    return out

if __name__ == "__main__":
    if len(sys.argv) < 2:
        raise SystemExit("Usage : python3 render.py <spec.json> [sortie.docx]")
    with open(sys.argv[1], encoding="utf-8") as f:
        spec = json.load(f)
    out = sys.argv[2] if len(sys.argv) > 2 else None
    print("Document genere :", render(spec, out))
